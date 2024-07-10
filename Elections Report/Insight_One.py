import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from tabulate import tabulate
from docx import Document
from docx.shared import Inches

# Load your data
df = pd.read_csv("CSV Files/All_States_PC_electionresults.csv", encoding="unicode_escape")
df.dropna(inplace=True)

# Convert columns to numeric
df['Postal Votes'] = pd.to_numeric(df['Postal Votes'], errors='coerce').fillna(0)
df['Total Votes'] = pd.to_numeric(df['Total Votes'], errors='coerce')

# Determine the winning parties
winning_parties = df.groupby('State').apply(lambda x: x.loc[x['Total Votes'].idxmax()])[['State', 'Party']]
winning_parties.reset_index(drop=True, inplace=True)

# Displaying as a table
print(tabulate(winning_parties, headers='keys', tablefmt='fancy_grid'))

# Create a Word document
doc = Document()



doc.add_heading('INSIGHTS OF LOK SABHA ELECTIONS - 2024', level=1)


doc.add_heading('KEY INSIGHT 1: National Party Dominance', level=2)
doc.add_paragraph(
    "The Bharatiya Janata Party (BJP) and the Indian National Congress (INC) continue to dominate the political landscape, securing the majority of votes across multiple states. "
    "This trend indicates a strong preference for national parties over regional parties.\n\n"
    "Let’s see the winning party across each state and the party that dominated the most across India.\n\n"
    "The link to the Winning Party List is provided here: Winning parties by state"
)


doc.add_heading('Winning Parties by State', level=1)

# TABLE OF INSIGHT ONE
table = doc.add_table(rows=1, cols=len(winning_parties.columns))
hdr_cells = table.rows[0].cells
for i, column_name in enumerate(winning_parties.columns):
    hdr_cells[i].text = column_name
for index, row in winning_parties.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

# Generate a pie chart
plt.figure(figsize=(8, 6))
party_counts = winning_parties['Party'].value_counts()
labels = party_counts.index
sizes = party_counts.values
plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
plt.axis('equal')  # Ensure the pie chart is drawn as a circle
plt.tight_layout()


plt.savefig('REPORT FILES/Winning_Parties_PieChart.png')

doc.add_paragraph()
doc.add_heading('Distribution of Votes among Winning Parties', level=2)
doc.add_picture('Winning_Parties_PieChart.png', width=Inches(6.0))

# INSIGHT 2 STARTS
doc.add_heading('KEY INSIGHT 2: Regional Party Influence', level=2)

doc.add_paragraph(
    "While national parties dominate, several regional parties have shown significant influence in specific states"
    "For instance, parties like the Trinamool Congress in West Bengal,Voice of the People Party in Meghalaya and Sikkim Krantikari Morcha in Sikkim have secured a substantial share of votes.\n\n"
    "Here are the list of the States where Regional Parties Dominated\n\n"
)


regional_parties = ['All India Trinamool Congress', 'Shiv Sena', 'DMK','Sikkim Krantikari Morcha','Zoram Peopleâs Movement','Voice of the People Party','Telugu Desam']  

# Filter out regional parties and tabulate
filtered_parties = winning_parties[winning_parties['Party'].isin(regional_parties)]

# Display as a table
print(tabulate(filtered_parties, headers='keys', tablefmt='fancy_grid'))

Regional_table = doc.add_table(rows=1, cols=len(filtered_parties.columns))
hdr_cells = Regional_table.rows[0].cells
for i, column_name in enumerate(filtered_parties.columns):
    hdr_cells[i].text = column_name
for index, row in filtered_parties.iterrows():
    row_cells = Regional_table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)


# Save the document
doc.save('REPORT FILES/Winning_Parties_Report.docx')

# Display the pie chart
plt.show()
