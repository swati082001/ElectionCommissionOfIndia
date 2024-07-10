import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
from tabulate import tabulate

# Load the CSV file
file_path = 'CSV Files/All_States_PC_electionresults.csv'
data = pd.read_csv(file_path)



# Convert relevant columns to numeric
data['EVM Votes'] = pd.to_numeric(data['EVM Votes'].str.replace(',', ''), errors='coerce')
data['Postal Votes'] = pd.to_numeric(data['Postal Votes'].str.replace(',', ''), errors='coerce')
data['Total Votes'] = pd.to_numeric(data['Total Votes'].str.replace(',', ''), errors='coerce')
data['% of Votes'] = pd.to_numeric(data['% of Votes'].str.replace('%', ''), errors='coerce')

# Setting the aesthetic style of the plots
sns.set(style="whitegrid")

# 1. Top Parties by Total Votes
top_parties = data.groupby('Party')['Total Votes'].sum().sort_values(ascending=False).head(10)
plt.figure(figsize=(12, 8))
sns.barplot(x=top_parties.values, y=top_parties.index)
plt.xlabel('Total Votes')
plt.ylabel('Party')
plt.savefig('REPORT FILES/Insight1.png')
# plt.show()

# 2. Candidates with Highest Votes
top_candidates = data[['Candidate', 'Total Votes', 'Party','State']].sort_values(by='Total Votes', ascending=False).head(10)
plt.figure(figsize=(12, 8))
sns.barplot(x='Total Votes', y='Candidate', hue='Party', dodge=False, data=top_candidates)
plt.xlabel('Total Votes')
plt.ylabel('Candidate')
plt.legend(title='Party')
plt.savefig('REPORT FILES/Insight2.png')
# plt.show()


# 3. Constituencies with Highest Voter Turnout
top_constituencies = data.groupby('Constituency Name')['Total Votes'].sum().sort_values(ascending=False).head(10)
plt.figure(figsize=(10, 6))
sns.barplot(x=top_constituencies.values, y=top_constituencies.index)
plt.xlabel('Total Votes')
plt.ylabel('Constituency Name')
plt.savefig('REPORT FILES/Insight3.png')
# plt.show()

# 4. State-wise Distribution of Votes
state_votes = data.groupby('State')['Total Votes'].sum().sort_values(ascending=False)
plt.figure(figsize=(12, 8))
sns.barplot(x=state_votes.values, y=state_votes.index)
plt.xlabel('Total Votes')
plt.ylabel('State')
plt.savefig('REPORT FILES/Insight4.png')
# plt.show()

# 5. Winning Margins
data['Winning Margin'] = data.groupby('Constituency Name')['Total Votes'].transform(lambda x: x.max() - x.nlargest(2).min())
winning_margins = data.sort_values(by='Winning Margin', ascending=False).head(10)
winning_margins_table = winning_margins[['Constituency Name', 'Candidate', 'Total Votes', 'Winning Margin']]

# Tabulate the results
print(tabulate(winning_margins_table, headers='keys', tablefmt='grid'))
plt.figure(figsize=(10, 6))
sns.barplot(x=winning_margins['Winning Margin'], y=winning_margins['Candidate'])
plt.xlabel('Winning Margin')
plt.ylabel('Candidate')
plt.savefig('REPORT FILES/Insight5.png')
# plt.show()

# 6. Correlation between EVM and Postal Votes
plt.figure(figsize=(6, 6))
sns.heatmap(data[['EVM Votes', 'Postal Votes']].corr(), annot=True, cmap='coolwarm')
plt.savefig('REPORT FILES/Insight6.png')
# plt.show()

# 7. Distribution of Vote Percentages
plt.figure(figsize=(10, 6))
sns.histplot(data['% of Votes'], bins=50, kde=True)
plt.xlabel('% of Votes')
plt.ylabel('Frequency')
plt.savefig('REPORT FILES/Insight7.png')
# plt.show()


# 8. Top States by Number of Candidates
top_states_candidates = data['State'].value_counts().head(10)
plt.figure(figsize=(10, 6))
sns.barplot(x=top_states_candidates.values, y=top_states_candidates.index)
plt.xlabel('Number of Candidates')
plt.ylabel('State')
plt.savefig('REPORT FILES/Insight8.png')
# plt.show()


# 9. Party Performance in Different States
party_state_performance = data.groupby(['State', 'Party'])['Total Votes'].sum().unstack().fillna(0)
plt.figure(figsize=(20, 10))
sns.heatmap(party_state_performance, cmap='viridis')
plt.xlabel('Party')
plt.ylabel('State')
plt.savefig('REPORT FILES/Insight9.png')
# plt.show()

# 10. Analysis of Independents
independents_performance = data[data['Party'] == 'Independent'].groupby('State')['Total Votes'].sum().sort_values(ascending=False).head(10)
plt.figure(figsize=(10, 6))
sns.barplot(x=independents_performance.values, y=independents_performance.index)
plt.xlabel('Total Votes')
plt.ylabel('State')
plt.savefig('REPORT FILES/Insight10.png')
# plt.show()




# # Save insights to a Word document
doc = Document()
# adding content to the document
doc.add_heading('INSIGHTS OF LOK SABHA ELECTIONS - 2024', level=1)


doc.add_heading('1. Top Parties by Total Votes', level=2)
doc.add_paragraph(
    "This insight reveals the political parties that received the highest total votes across all constituencies. It helps to identify the most popular and influential parties in the election. The data shows the sum of votes for each party and ranks them to highlight the top 10 parties."
    "The Bharatiya Janata Party (BJP) and the Indian National Congress (INC) continue to dominate the political landscape, securing the majority of votes across multiple states. "
    "This trend indicates a strong preference for national parties over regional parties.\n\n"
    "Let’s see the winning party across each state and the party that dominated the most across India.\n\n"
    "The link to the Winning Party List is provided here: Winning parties by state"
)



doc.add_heading('Top Parties by Total Votes', level=2)
doc.add_picture('REPORT FILES/Insight1.png', width=Inches(7.0))

# INSIGHT 2 STARTS
doc.add_heading('2. Candidates with Highest Votes', level=2)

doc.add_paragraph(
    "This insight lists the candidates who received the highest number of votes in the election. It highlights the top-performing candidates and their respective parties. The data shows the sum of votes for each member and ranks them to highlight the top 10 members across India.\n"
    "Even though the Bharatiya Janata Party (BJP) won with the highest number of votes as a party, the candidate who received the highest number of votes was RAKIBUL HUSSAIN from the Indian National Congress(INC)."
)


doc.add_heading('Candidates with Highest Votes', level=2)
doc.add_picture('REPORT FILES/Insight2.png', width=Inches(7.0))


# INSIGHT 3 STARTS
doc.add_heading('3. Constituencies with Highest Voter Turnout', level=2)

doc.add_paragraph(
    "This insight identifies the constituencies that had the highest voter turnout by summing the total votes cast in each constituency and ranking them.\n"
    "An analysis reveals that a significant number of the top 10 constituencies with the highest voter turnout are from the state of Assam. This indicates a stronger voter engagement. Such engagement can be attributed to effective voter mobilization efforts, a high level of political awareness, and socio-economic factors that motivate voters to participate actively."
)


doc.add_heading('Constituencies with Highest Voter Turnout', level=2)
doc.add_picture('REPORT FILES/Insight3.png', width=Inches(7.0))


# INSIGHT 4 STARTS
doc.add_heading('4. State-wise Distribution of Votes', level=2)

doc.add_paragraph(
    "This insight provides a state-wise distribution of total votes, showing the sum of votes for each state. It helps to understand the geographical spread of voter support across different regions in India.\n"
    "By examining this data, we observe that Uttar Pradesh (UP) has the highest total votes cast, winning by a clear margin. This significant lead can be attributed to UP's large population, which is the highest among all Indian states, thereby resulting in a higher number of eligible voters and, consequently, more votes cast.\n"
    "Conversely, states like Andaman & Nicobar Islands and Lakshadweep have the least vote counts. Their low total votes can be attributed to their smaller populations and geographic locations, which limit the number of eligible voters."
)


doc.add_heading('State-wise Distribution of Votes', level=2)
doc.add_picture('REPORT FILES/Insight4.png', width=Inches(7.0))


# INSIGHT 5 STARTS
doc.add_heading('5. Winning Margins', level=2)

doc.add_paragraph(
    "This insight calculates the winning margins for candidates, showing the difference between the votes received by the winning candidate and the runner-up. It lists the top 10 candidates with the highest winning margins.\n"
    "The data reveals significant victories in constituencies such as Dhubri, where the winning margins are substantial, indicating strong support for the winning candidates. High winning margins can indicate a strong voter preference for the leading candidates and effective campaigning strategies."
)


doc.add_heading('Winning Margins', level=2)
table = doc.add_table(rows=1, cols=len(winning_margins_table.columns))
hdr_cells = table.rows[0].cells

for i, column in enumerate(winning_margins_table.columns):
    hdr_cells[i].text = column

# Add the data rows
for index, row in winning_margins_table.iterrows():
    row_cells = table.add_row().cells
    for i, value in enumerate(row):
        row_cells[i].text = str(value)

doc.add_picture('REPORT FILES/Insight5.png', width=Inches(7.0))


# INSIGHT 6 STARTS
doc.add_heading('6. Correlation between EVM and Postal Votes', level=2)

doc.add_paragraph(
    "This insight examines the correlation between votes received through Electronic Voting Machines (EVMs) and postal votes. A correlation coefficient is calculated to show the relationship between these two variables. Understanding this relationship helps to determine if there is consistency in voter preferences across different voting methods."
    "This insight examines the correlation between votes received through Electronic Voting Machines (EVMs) and postal votes. A correlation coefficient is calculated to show the relationship between these two variables. Understanding this relationship helps to determine if there is consistency in voter preferences across different voting methods."
    "The strong correlation might be due to similar factors influencing voter behavior across both voting methods, such as party loyalty, candidate popularity, or socio-political factors within constituencies."

)


doc.add_heading('Correlation between EVM and Postal Votes', level=2)
doc.add_picture('REPORT FILES/Insight6.png', width=Inches(7.0))


# INSIGHT 7 STARTS
doc.add_heading('7. Distribution of Vote Percentages', level=2)

doc.add_paragraph(
    "This insight provides a statistical summary of the distribution of vote percentages among candidates. It includes metrics like mean, median, and standard deviation. This metric helps in understanding the overall competitiveness of the election. It shows how vote percentages are distributed among candidates, indicating whether the votes are concentrated among a few candidates or spread out."
    "This metric helps in understanding the overall competitiveness of the election.It shows how vote percentages are distributed among candidates, indicating whether the votes are concentrated among a few candidates or spread out."

)


doc.add_heading('Distribution of Vote Percentages', level=2)
doc.add_picture('REPORT FILES/Insight7.png', width=Inches(7.0))


# INSIGHT 8 STARTS
doc.add_heading('8. Top States by Number of Candidates', level=2)

doc.add_paragraph(
    "This insight lists the states with the highest number of candidates contesting in the election. It shows the count of candidates for each state and ranks them to highlight the top 10 states."
    "This metric helps in understanding the level of political competition in different states. A higher number of candidates might indicate more political diversity and competition. Maharashtra, Tamil Nadu, and Uttar Pradesh suggest significant political diversity and competition. These states likely have a diverse range of political parties and independent candidates, making the election highly competitive."

)


doc.add_heading('Top States by Number of Candidates', level=2)
doc.add_picture('REPORT FILES/Insight8.png', width=Inches(7.0))


# INSIGHT 9 STARTS
doc.add_heading('9. Party Performance in Different States', level=2)

doc.add_paragraph(
    "This insight provides a detailed view of how different political parties performed across various states. It shows the total votes received by each party in each state. This metric is crucial for understanding the regional strengths and weaknesses of different parties. It helps in identifying which parties are dominant in specific states."

)


doc.add_heading('Party Performance in Different States', level=2)
doc.add_picture('REPORT FILES/Insight9.png', width=Inches(7.0))


# INSIGHT 10 STARTS
doc.add_heading('10. Analysis of Independents', level=2)

doc.add_paragraph(
    "This insight focuses on the performance of independent candidates. It lists the top states where independent candidates received the highest total votes.This metric helps in understanding the support for non-affiliated candidates.It shows regions where voters prefer independent candidates over party-affiliated ones. Maharashtra, Bihar and, Jharkhand suggest significant number of Independent candidates."

)


doc.add_heading('Analysis of Independents', level=2)
doc.add_picture('REPORT FILES/Insight10.png', width=Inches(7.0))



doc.save('REPORT FILES/Election_commission_Report.docx')
